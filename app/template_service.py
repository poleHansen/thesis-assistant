from __future__ import annotations

import json
import re
from pathlib import Path

from app.domain import ProjectCreate, TemplateManifest, TemplateSource
from app.template_library import TEMPLATE_LIBRARY, TemplateBundle


DEFAULT_USER_SECTIONS = [
    "封面",
    "摘要",
    "Abstract",
    "目录",
    "引言",
    "相关工作",
    "方法",
    "实验",
    "结论",
    "参考文献",
]

DEFAULT_USER_STYLE_MAPPING = {
    "title": "Title",
    "chapter": "Heading 1",
    "section": "Heading 2",
    "body": "Normal",
    "caption": "Caption",
}

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([\w\-.\u4e00-\u9fff]+)\s*\}\}")


class TemplateService:
    def __init__(self, template_library: dict[str, TemplateBundle] | None = None) -> None:
        self.template_library = template_library or TEMPLATE_LIBRARY

    def choose_default_template(
        self, request: ProjectCreate
    ) -> tuple[TemplateSource, TemplateManifest]:
        if "课程" in request.school_requirements or "报告" in request.school_requirements:
            key = "course_project_report"
        elif request.paper_type == "algorithm":
            key = "engineering_thesis"
        else:
            key = "general_undergraduate"

        bundle = self.template_library[key]
        manifest = self._load_manifest(bundle)
        manifest = self._merge_word_template_hints(bundle.word_template_path, manifest)
        source = TemplateSource(
            source_type="library_default",
            template_id=bundle.template_id,
            template_name=bundle.template_name,
            template_path=str(bundle.word_template_path),
            ppt_template_path=str(bundle.ppt_template_path),
        )
        return source, manifest

    def parse_user_template(
        self, template_path: Path
    ) -> tuple[TemplateSource, TemplateManifest]:
        section_mapping, style_mapping, cover_fields = self._extract_docx_hints(
            template_path,
            DEFAULT_USER_SECTIONS,
            DEFAULT_USER_STYLE_MAPPING,
            ["学校", "学院", "专业", "题目", "作者", "学号", "指导教师", "日期"],
        )

        source = TemplateSource(
            source_type="user_upload",
            template_id=template_path.stem,
            template_name=template_path.name,
            template_path=str(template_path),
        )
        manifest = TemplateManifest(
            section_mapping=section_mapping,
            style_mapping=style_mapping,
            cover_fields=cover_fields,
            figure_slots=["图1", "图2", "图3"],
            table_slots=["表1", "表2", "表3"],
            citation_style="GB/T 7714",
            header_footer_rules={"header": "按用户模板输出", "footer": "自动分页"},
            toc_rules={"enabled": True, "depth": 3},
            ppt_layouts=["title", "content", "chart", "summary"],
        )
        return source, manifest

    def _load_manifest(self, bundle: TemplateBundle) -> TemplateManifest:
        if not bundle.manifest_path.exists():
            raise FileNotFoundError(f"Template manifest not found: {bundle.manifest_path}")

        try:
            payload = json.loads(bundle.manifest_path.read_text(encoding="utf-8"))
            return TemplateManifest(**payload)
        except (json.JSONDecodeError, TypeError, ValueError) as exc:
            raise ValueError(f"Invalid template manifest: {bundle.manifest_path}") from exc

    def _merge_word_template_hints(
        self,
        template_path: Path,
        manifest: TemplateManifest,
    ) -> TemplateManifest:
        sections, styles, cover_fields = self._extract_docx_hints(
            template_path,
            manifest.section_mapping,
            manifest.style_mapping,
            manifest.cover_fields,
        )
        return TemplateManifest(
            section_mapping=sections,
            style_mapping=styles,
            cover_fields=cover_fields,
            figure_slots=manifest.figure_slots,
            table_slots=manifest.table_slots,
            citation_style=manifest.citation_style,
            header_footer_rules=manifest.header_footer_rules,
            toc_rules=manifest.toc_rules,
            ppt_layouts=manifest.ppt_layouts,
        )

    def _extract_docx_hints(
        self,
        template_path: Path,
        fallback_sections: list[str],
        fallback_styles: dict[str, str],
        fallback_cover_fields: list[str],
    ) -> tuple[list[str], dict[str, str], list[str]]:
        section_mapping = list(fallback_sections)
        style_mapping = dict(fallback_styles)
        cover_fields = list(fallback_cover_fields)

        try:
            from docx import Document  # type: ignore

            document = Document(str(template_path))
            headings: list[str] = []
            styles_seen: dict[str, str] = {}
            section_placeholders: list[str] = []
            cover_placeholders: list[str] = []
            for paragraph in self._iter_docx_paragraphs(document):
                text = paragraph.text.strip()
                style_name = getattr(paragraph.style, "name", "") or "Normal"
                if style_name.startswith("Heading") and text:
                    headings.append(text)
                self._collect_placeholder_names(
                    text,
                    section_placeholders=section_placeholders,
                    cover_placeholders=cover_placeholders,
                )
                if "title" in style_name.lower():
                    styles_seen["title"] = style_name
                elif style_name.startswith("Heading 1"):
                    styles_seen["chapter"] = style_name
                elif style_name.startswith("Heading 2"):
                    styles_seen["section"] = style_name
                elif "caption" in style_name.lower():
                    styles_seen["caption"] = style_name
            if section_placeholders:
                section_mapping = section_placeholders[:20]
            elif headings:
                section_mapping = headings[:20]
            if cover_placeholders:
                cover_fields = cover_placeholders
            style_mapping.update(styles_seen)
        except Exception:
            pass

        return section_mapping, style_mapping, cover_fields

    def _iter_docx_paragraphs(self, document):
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

    def _collect_placeholder_names(
        self,
        text: str,
        *,
        section_placeholders: list[str],
        cover_placeholders: list[str],
    ) -> None:
        for match in PLACEHOLDER_PATTERN.finditer(text):
            placeholder = match.group(1)
            if placeholder.startswith("section."):
                section_name = placeholder.split(".", 1)[1].strip()
                if section_name and section_name not in section_placeholders:
                    section_placeholders.append(section_name)
            elif placeholder.startswith("cover."):
                field_name = placeholder.split(".", 1)[1].strip()
                if field_name and field_name not in cover_placeholders:
                    cover_placeholders.append(field_name)
