from __future__ import annotations

import shutil
import unittest
import uuid
from pathlib import Path

from app.domain import ProjectCreate
from app.template_library import build_template_library
from app.template_service import TemplateService

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - environment-dependent
    Document = None


class TemplateServiceTest(unittest.TestCase):
    def test_choose_default_template_for_algorithm_reads_manifest_from_library(self) -> None:
        root = Path("tests_runtime") / f"template-library-{uuid.uuid4().hex[:8]}"
        manifest_dir = root / "engineering_thesis"
        manifest_dir.mkdir(parents=True, exist_ok=True)
        try:
            (manifest_dir / "manifest.json").write_text(
                (
                    '{\n'
                    '  "section_mapping": ["封面", "第3章 方法设计"],\n'
                    '  "style_mapping": {"title": "Title", "chapter": "Heading 1", "section": "Heading 2", "body": "Normal"},\n'
                    '  "cover_fields": ["学校"],\n'
                    '  "figure_slots": ["图1"],\n'
                    '  "table_slots": ["表1"],\n'
                    '  "citation_style": "GB/T 7714",\n'
                    '  "header_footer_rules": {"header": "工科毕业设计", "footer": "第 X 页"},\n'
                    '  "toc_rules": {"enabled": true, "depth": 3},\n'
                    '  "ppt_layouts": ["title", "summary"]\n'
                    '}'
                ),
                encoding="utf-8",
            )
            service = TemplateService(build_template_library(root))

            source, manifest = service.choose_default_template(ProjectCreate(topic="图像分类算法"))
        finally:
            shutil.rmtree(root, ignore_errors=True)

        self.assertEqual(source.template_id, "engineering_thesis")
        self.assertTrue(source.template_path.endswith("engineering_thesis\\word\\template.docx"))
        self.assertTrue(source.ppt_template_path.endswith("engineering_thesis\\ppt\\template.pptx"))
        self.assertIn("第3章 方法设计", manifest.section_mapping)

    def test_parse_user_template_falls_back_without_valid_docx(self) -> None:
        service = TemplateService()
        temp_dir = Path("tests_runtime") / f"template-{uuid.uuid4().hex[:8]}"
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            fake_path = temp_dir / "school-template.docx"
            fake_path.write_bytes(b"not-a-real-docx")
            source, manifest = service.parse_user_template(fake_path)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
        self.assertEqual(source.source_type, "user_upload")
        self.assertEqual(source.template_name, "school-template.docx")
        self.assertTrue(manifest.section_mapping)

    @unittest.skipIf(Document is None, "python-docx is not installed in the current environment")
    def test_parse_user_template_prefers_placeholder_sections(self) -> None:
        service = TemplateService()
        temp_dir = Path("tests_runtime") / f"template-{uuid.uuid4().hex[:8]}"
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            template_path = temp_dir / "school-template.docx"
            document = Document()
            document.add_paragraph("{{cover.题目}}", style="Title")
            document.add_heading("学校自定义标题", level=1)
            document.add_paragraph("{{section.摘要}}")
            document.add_paragraph("{{section.第1章 绪论}}")
            document.save(template_path)

            source, manifest = service.parse_user_template(template_path)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

        self.assertEqual(source.template_name, "school-template.docx")
        self.assertEqual(manifest.cover_fields[0], "题目")
        self.assertEqual(manifest.section_mapping[:2], ["摘要", "第1章 绪论"])


if __name__ == "__main__":
    unittest.main()
